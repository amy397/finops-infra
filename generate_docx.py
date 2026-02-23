#!/usr/bin/env python3
"""
FinOps Infrastructure Terraform 코드 산출물 문서 생성 스크립트
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """셀 텍스트 서식 설정"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = color


def add_header_row(table, row_idx, texts, bg_color="2E4057"):
    """테이블 헤더 행 설정"""
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        set_cell_shading(row.cells[i], bg_color)
        set_cell_text(row.cells[i], text, bold=True, size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      color=RGBColor(0xFF, 0xFF, 0xFF))


def add_data_row(table, row_idx, texts, alignments=None):
    """테이블 데이터 행 설정"""
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        align = WD_ALIGN_PARAGRAPH.LEFT
        if alignments and i < len(alignments):
            align = alignments[i]
        set_cell_text(row.cells[i], text, size=9, alignment=align)


def read_tf_file(relative_path):
    """Terraform 파일 읽기"""
    full_path = os.path.join(BASE_DIR, relative_path)
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read()


def add_code_block(doc, code_text, title=None):
    """코드 블록 추가"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(9)
        run.bold = True
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # 코드를 테이블로 감싸서 배경색 표현
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")

    cell.text = ""
    for line in code_text.strip().split("\n"):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(7.5)
        run.font.name = "Consolas"
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)

    # 첫 빈 paragraph 제거
    if cell.paragraphs[0].text == "":
        p_element = cell.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    # 테이블 테두리 설정
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # 테이블 너비를 페이지 전체로 설정
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tbl_width)


def add_heading_styled(doc, text, level=1):
    """스타일이 적용된 제목 추가"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return heading


def create_document():
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── 기본 스타일 설정 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ════════════════════════════════════════
    # 표지
    # ════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("FinOps 인프라 코드 산출물")
    run.font.size = Pt(28)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Terraform Infrastructure as Code")
    run.font.size = Pt(16)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 표지 정보 테이블
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("프로젝트명", "FinOps Infrastructure"),
        ("클라우드 플랫폼", "AWS (Amazon Web Services)"),
        ("리전", "ap-northeast-2 (서울)"),
        ("IaC 도구", "Terraform >= 1.0.0"),
        ("작성일", "2026-02-23"),
    ]
    for i, (label, value) in enumerate(info_data):
        set_cell_shading(info_table.rows[i].cells[0], "2E4057")
        set_cell_text(info_table.rows[i].cells[0], label, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_text(info_table.rows[i].cells[1], value, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 페이지 나누기
    doc.add_page_break()

    # ════════════════════════════════════════
    # 목차
    # ════════════════════════════════════════
    add_heading_styled(doc, "목차", level=1)
    toc_items = [
        "1. 개요",
        "    1.1 프로젝트 개요",
        "    1.2 시스템 아키텍처",
        "    1.3 디렉토리 구조",
        "2. 모듈 상세",
        "    2.1 VPC 모듈",
        "    2.2 Security Group 모듈",
        "    2.3 EC2 모듈",
        "    2.4 ECR 모듈",
        "    2.5 RDS 모듈",
        "3. 환경 구성",
        "    3.1 Dev 환경",
        "4. 변수 명세",
        "5. 출력값 명세",
        "6. 리소스 요약",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. 개요
    # ════════════════════════════════════════
    add_heading_styled(doc, "1. 개요", level=1)

    # 1.1 프로젝트 개요
    add_heading_styled(doc, "1.1 프로젝트 개요", level=2)
    overview_text = (
        "본 문서는 FinOps(Financial Operations) 인프라스트럭처를 코드로 정의한 "
        "Terraform 프로젝트의 코드 산출물 문서입니다. "
        "이 프로젝트는 AWS 클라우드 환경에서 비용 최적화 모니터링 및 알림 시스템을 위한 "
        "인프라를 자동으로 프로비저닝합니다."
    )
    p = doc.add_paragraph(overview_text)
    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # 프로젝트 기본 정보 테이블
    proj_table = doc.add_table(rows=7, cols=2)
    proj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(proj_table, 0, ["항목", "내용"])
    proj_data = [
        ("프로젝트명", "FinOps Infrastructure"),
        ("IaC 도구", "Terraform >= 1.0.0"),
        ("클라우드 프로바이더", "AWS (hashicorp/aws ~> 5.0)"),
        ("배포 리전", "ap-northeast-2 (서울)"),
        ("환경", "dev (개발)"),
        ("마이크로서비스 수", "5개 (cost-api, cost-collector, alert-service, gateway, dashboard)"),
    ]
    for i, (label, value) in enumerate(proj_data):
        add_data_row(proj_table, i + 1, [label, value])

    doc.add_paragraph()

    # 1.2 시스템 아키텍처
    add_heading_styled(doc, "1.2 시스템 아키텍처", level=2)
    arch_desc = (
        "본 인프라는 모듈화된 Terraform 구성을 통해 VPC, Security Group, EC2, ECR, RDS 등의 "
        "AWS 리소스를 체계적으로 관리합니다. 퍼블릭 서브넷에 EC2 인스턴스를 배치하고, "
        "프라이빗 서브넷에 RDS 데이터베이스를 격리하여 보안성을 확보합니다."
    )
    doc.add_paragraph(arch_desc)

    # 아키텍처 다이어그램 (텍스트)
    arch_diagram = """
+---------------------------------------------------------------+
|                    VPC (10.0.0.0/16)                          |
|                                                               |
|   +----------------------------+                              |
|   |  Public Subnet             |     Internet Gateway         |
|   |  10.0.1.0/24 (AZ-2a)      |<------->  IGW               |
|   |                            |                              |
|   |  +----------------------+  |                              |
|   |  |  EC2 (t3.micro)     |  |     +-------------------+   |
|   |  |  - Docker            |  |     | ECR Registry      |   |
|   |  |  - Docker Compose    |  |---->| - cost-api        |   |
|   |  |  - Git               |  |     | - cost-collector  |   |
|   |  |  - IAM Role (ECR)   |  |     | - alert-service   |   |
|   |  +----------------------+  |     | - gateway         |   |
|   +----------------------------+     | - dashboard       |   |
|              |                       +-------------------+   |
|              | Port 5432 (SG restricted)                     |
|              v                                               |
|   +-------------------------------------------+             |
|   |  Private Subnets                          |             |
|   |  10.0.10.0/24 (AZ-2a)                    |             |
|   |  10.0.20.0/24 (AZ-2c)                    |             |
|   |                                           |             |
|   |  +-------------------------------------+  |             |
|   |  |  RDS PostgreSQL 15 (db.t3.micro)   |  |             |
|   |  |  DB: finops / User: postgres        |  |             |
|   |  |  Storage: 20GB gp2                  |  |             |
|   |  +-------------------------------------+  |             |
|   +-------------------------------------------+             |
+---------------------------------------------------------------+
"""
    add_code_block(doc, arch_diagram, title="[시스템 아키텍처 다이어그램]")
    doc.add_paragraph()

    # 1.3 디렉토리 구조
    add_heading_styled(doc, "1.3 디렉토리 구조", level=2)
    dir_structure = """finops-infra/
├── modules/                          # 재사용 가능한 Terraform 모듈
│   ├── vpc/                          # VPC 네트워크 모듈
│   │   ├── main.tf                   # VPC, 서브넷, IGW, 라우트 테이블
│   │   ├── variables.tf              # 입력 변수 정의
│   │   └── outputs.tf                # 출력값 정의
│   ├── security-group/               # 보안 그룹 모듈
│   │   ├── main.tf                   # EC2/RDS 보안 그룹 규칙
│   │   ├── variables.tf              # 입력 변수 정의
│   │   └── outputs.tf                # 출력값 정의
│   ├── ec2/                          # EC2 인스턴스 모듈
│   │   ├── main.tf                   # AMI, IAM, EC2 인스턴스
│   │   ├── variables.tf              # 입력 변수 정의
│   │   └── outputs.tf                # 출력값 정의
│   ├── ecr/                          # ECR 컨테이너 레지스트리 모듈
│   │   ├── main.tf                   # ECR 리포지토리, 수명주기 정책
│   │   ├── variables.tf              # 입력 변수 정의
│   │   └── output.tf                 # 출력값 정의
│   └── rds/                          # RDS 데이터베이스 모듈
│       ├── main.tf                   # DB 서브넷 그룹, RDS 인스턴스
│       ├── variables.tf              # 입력 변수 정의
│       └── outputs.tf                # 출력값 정의
├── environments/                     # 환경별 구성
│   └── dev/                          # 개발 환경
│       ├── main.tf                   # 모듈 오케스트레이션
│       ├── variables.tf              # 환경 변수 정의
│       └── output.tf                 # 환경 출력값
├── .gitignore                        # Git 제외 파일 목록
└── generate_docx.py                  # 본 문서 생성 스크립트"""
    add_code_block(doc, dir_structure)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. 모듈 상세
    # ════════════════════════════════════════
    add_heading_styled(doc, "2. 모듈 상세", level=1)

    # --- 2.1 VPC 모듈 ---
    add_heading_styled(doc, "2.1 VPC 모듈", level=2)
    doc.add_paragraph(
        "VPC(Virtual Private Cloud) 모듈은 AWS 네트워크 인프라의 기반을 구성합니다. "
        "VPC, 인터넷 게이트웨이, 퍼블릭/프라이빗 서브넷, 라우트 테이블을 생성합니다."
    )
    doc.add_paragraph()

    # VPC 리소스 목록 테이블
    vpc_res_table = doc.add_table(rows=8, cols=4)
    vpc_res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(vpc_res_table, 0, ["No", "리소스 타입", "리소스 이름", "설명"])
    vpc_resources = [
        ("1", "aws_vpc", "main", "VPC 생성 (CIDR: 10.0.0.0/16)"),
        ("2", "aws_internet_gateway", "main", "인터넷 게이트웨이"),
        ("3", "aws_subnet", "public", "퍼블릭 서브넷 (10.0.1.0/24, AZ-2a)"),
        ("4", "aws_subnet", "private", "프라이빗 서브넷 (10.0.10.0/24, AZ-2a)"),
        ("5", "aws_subnet", "private_2", "프라이빗 서브넷 2 (10.0.20.0/24, AZ-2c)"),
        ("6", "aws_route_table", "public", "퍼블릭 라우트 테이블 (0.0.0.0/0 -> IGW)"),
        ("7", "aws_route_table_association", "public", "퍼블릭 서브넷-라우트 테이블 연결"),
    ]
    for i, data in enumerate(vpc_resources):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(vpc_res_table, i + 1, list(data),
                     [center, WD_ALIGN_PARAGRAPH.LEFT, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()
    add_heading_styled(doc, "modules/vpc/main.tf", level=3)
    vpc_main = read_tf_file("modules/vpc/main.tf")
    add_code_block(doc, vpc_main)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/vpc/variables.tf", level=3)
    vpc_vars = read_tf_file("modules/vpc/variables.tf")
    add_code_block(doc, vpc_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/vpc/outputs.tf", level=3)
    vpc_outputs = read_tf_file("modules/vpc/outputs.tf")
    add_code_block(doc, vpc_outputs)

    doc.add_page_break()

    # --- 2.2 Security Group 모듈 ---
    add_heading_styled(doc, "2.2 Security Group 모듈", level=2)
    doc.add_paragraph(
        "보안 그룹 모듈은 EC2 인스턴스와 RDS 데이터베이스에 대한 네트워크 접근 제어를 정의합니다. "
        "EC2는 외부 트래픽을 허용하고, RDS는 EC2에서만 접근 가능하도록 제한합니다."
    )
    doc.add_paragraph()

    # EC2 보안 그룹 인바운드 규칙 테이블
    p = doc.add_paragraph()
    run = p.add_run("[EC2 보안 그룹 인바운드 규칙]")
    run.bold = True
    run.font.size = Pt(10)

    sg_ec2_table = doc.add_table(rows=6, cols=4)
    sg_ec2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(sg_ec2_table, 0, ["포트", "프로토콜", "소스", "용도"])
    sg_ec2_rules = [
        ("22", "TCP", "0.0.0.0/0", "SSH 원격 접속"),
        ("80", "TCP", "0.0.0.0/0", "HTTP 웹 트래픽"),
        ("443", "TCP", "0.0.0.0/0", "HTTPS 웹 트래픽"),
        ("3000", "TCP", "0.0.0.0/0", "Dashboard (React 프론트엔드)"),
        ("8282", "TCP", "0.0.0.0/0", "Gateway (API 게이트웨이)"),
    ]
    for i, data in enumerate(sg_ec2_rules):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(sg_ec2_table, i + 1, list(data), [center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # RDS 보안 그룹 인바운드 규칙 테이블
    p = doc.add_paragraph()
    run = p.add_run("[RDS 보안 그룹 인바운드 규칙]")
    run.bold = True
    run.font.size = Pt(10)

    sg_rds_table = doc.add_table(rows=2, cols=4)
    sg_rds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(sg_rds_table, 0, ["포트", "프로토콜", "소스", "용도"])
    add_data_row(sg_rds_table, 1, ["5432", "TCP", "EC2 보안 그룹", "PostgreSQL 접근"],
                 [WD_ALIGN_PARAGRAPH.CENTER] * 3 + [WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()
    add_heading_styled(doc, "modules/security-group/main.tf", level=3)
    sg_main = read_tf_file("modules/security-group/main.tf")
    add_code_block(doc, sg_main)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/security-group/variables.tf", level=3)
    sg_vars = read_tf_file("modules/security-group/variables.tf")
    add_code_block(doc, sg_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/security-group/outputs.tf", level=3)
    sg_outputs = read_tf_file("modules/security-group/outputs.tf")
    add_code_block(doc, sg_outputs)

    doc.add_page_break()

    # --- 2.3 EC2 모듈 ---
    add_heading_styled(doc, "2.3 EC2 모듈", level=2)
    doc.add_paragraph(
        "EC2 모듈은 애플리케이션 서버 인스턴스를 생성합니다. Amazon Linux 2023 기반으로 "
        "Docker, Docker Compose, Git이 자동 설치되며, ECR 읽기 권한이 부여된 IAM 역할이 연결됩니다."
    )
    doc.add_paragraph()

    # EC2 리소스 테이블
    ec2_res_table = doc.add_table(rows=6, cols=4)
    ec2_res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(ec2_res_table, 0, ["No", "리소스 타입", "리소스 이름", "설명"])
    ec2_resources = [
        ("1", "aws_ami (data)", "amazon_linux", "Amazon Linux 2023 최신 AMI 조회"),
        ("2", "aws_iam_role", "ec2_role", "EC2 IAM 역할 (ECR 접근용)"),
        ("3", "aws_iam_role_policy_attachment", "ecr_read", "ECR ReadOnly 정책 연결"),
        ("4", "aws_iam_instance_profile", "ec2_profile", "IAM 인스턴스 프로파일"),
        ("5", "aws_instance", "app", "EC2 인스턴스 (t3.micro)"),
    ]
    for i, data in enumerate(ec2_resources):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(ec2_res_table, i + 1, list(data),
                     [center, WD_ALIGN_PARAGRAPH.LEFT, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # User Data 스크립트 설명
    p = doc.add_paragraph()
    run = p.add_run("[EC2 User Data - 초기화 스크립트]")
    run.bold = True
    run.font.size = Pt(10)

    userdata_table = doc.add_table(rows=4, cols=2)
    userdata_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(userdata_table, 0, ["설치 항목", "설명"])
    userdata_items = [
        ("Docker", "컨테이너 런타임 설치 및 서비스 활성화"),
        ("Docker Compose", "멀티 컨테이너 오케스트레이션 도구"),
        ("Git", "소스 코드 관리 도구"),
    ]
    for i, (item, desc) in enumerate(userdata_items):
        add_data_row(userdata_table, i + 1, [item, desc])

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ec2/main.tf", level=3)
    ec2_main = read_tf_file("modules/ec2/main.tf")
    add_code_block(doc, ec2_main)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ec2/variables.tf", level=3)
    ec2_vars = read_tf_file("modules/ec2/variables.tf")
    add_code_block(doc, ec2_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ec2/outputs.tf", level=3)
    ec2_outputs = read_tf_file("modules/ec2/outputs.tf")
    add_code_block(doc, ec2_outputs)

    doc.add_page_break()

    # --- 2.4 ECR 모듈 ---
    add_heading_styled(doc, "2.4 ECR 모듈", level=2)
    doc.add_paragraph(
        "ECR(Elastic Container Registry) 모듈은 Docker 컨테이너 이미지를 저장하기 위한 "
        "프라이빗 레지스트리를 생성합니다. 5개의 마이크로서비스 각각에 대한 리포지토리를 생성하고, "
        "비용 절감을 위한 이미지 수명주기 정책을 적용합니다."
    )
    doc.add_paragraph()

    # ECR 리포지토리 목록
    ecr_repo_table = doc.add_table(rows=6, cols=3)
    ecr_repo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(ecr_repo_table, 0, ["No", "리포지토리명", "용도"])
    ecr_repos = [
        ("1", "cost-api", "비용 데이터 API 서비스"),
        ("2", "cost-collector", "비용 데이터 수집 서비스"),
        ("3", "alert-service", "알림/알람 서비스"),
        ("4", "gateway", "API 게이트웨이 서비스 (포트 8282)"),
        ("5", "dashboard", "React 프론트엔드 대시보드 (포트 3000)"),
    ]
    for i, data in enumerate(ecr_repos):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(ecr_repo_table, i + 1, list(data),
                     [center, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("[수명주기 정책]")
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph("  - 각 리포지토리에 최근 5개 이미지만 유지하는 정책을 적용하여 스토리지 비용을 절감합니다.")
    doc.add_paragraph("  - Push 시 자동 이미지 스캔이 활성화되어 보안 취약점을 점검합니다.")

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ecr/main.tf", level=3)
    ecr_main = read_tf_file("modules/ecr/main.tf")
    add_code_block(doc, ecr_main)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ecr/variables.tf", level=3)
    ecr_vars = read_tf_file("modules/ecr/variables.tf")
    add_code_block(doc, ecr_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/ecr/output.tf", level=3)
    ecr_outputs = read_tf_file("modules/ecr/output.tf")
    add_code_block(doc, ecr_outputs)

    doc.add_page_break()

    # --- 2.5 RDS 모듈 ---
    add_heading_styled(doc, "2.5 RDS 모듈", level=2)
    doc.add_paragraph(
        "RDS 모듈은 PostgreSQL 15 데이터베이스 인스턴스를 프라이빗 서브넷에 생성합니다. "
        "개발 환경에 최적화된 설정으로, 프리티어 호환 인스턴스를 사용합니다."
    )
    doc.add_paragraph()

    # RDS 구성 정보 테이블
    rds_config_table = doc.add_table(rows=11, cols=2)
    rds_config_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(rds_config_table, 0, ["구성 항목", "설정값"])
    rds_config = [
        ("엔진", "PostgreSQL 15"),
        ("인스턴스 클래스", "db.t3.micro (프리티어)"),
        ("스토리지", "20 GB (gp2)"),
        ("데이터베이스명", "finops"),
        ("사용자명", "postgres"),
        ("공개 접근", "비활성화 (프라이빗 서브넷)"),
        ("Multi-AZ", "비활성화 (개발 환경)"),
        ("암호화", "비활성화 (개발 환경)"),
        ("백업 보존 기간", "0일 (개발 환경)"),
        ("최종 스냅샷", "건너뛰기 (skip_final_snapshot)"),
    ]
    for i, (item, value) in enumerate(rds_config):
        add_data_row(rds_config_table, i + 1, [item, value])

    doc.add_paragraph()
    add_heading_styled(doc, "modules/rds/main.tf", level=3)
    rds_main = read_tf_file("modules/rds/main.tf")
    add_code_block(doc, rds_main)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/rds/variables.tf", level=3)
    rds_vars = read_tf_file("modules/rds/variables.tf")
    add_code_block(doc, rds_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "modules/rds/outputs.tf", level=3)
    rds_outputs = read_tf_file("modules/rds/outputs.tf")
    add_code_block(doc, rds_outputs)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 3. 환경 구성
    # ════════════════════════════════════════
    add_heading_styled(doc, "3. 환경 구성", level=1)
    add_heading_styled(doc, "3.1 Dev 환경", level=2)
    doc.add_paragraph(
        "개발(dev) 환경은 모든 모듈을 통합하여 전체 인프라를 오케스트레이션합니다. "
        "Terraform >= 1.0.0 및 AWS Provider ~> 5.0을 요구합니다."
    )
    doc.add_paragraph()

    # 모듈 의존성 테이블
    dep_table = doc.add_table(rows=6, cols=3)
    dep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(dep_table, 0, ["모듈", "소스 경로", "의존성"])
    dep_data = [
        ("VPC", "../../modules/vpc", "없음"),
        ("Security Group", "../../modules/security-group", "VPC (vpc_id)"),
        ("EC2", "../../modules/ec2", "VPC (subnet_id), SG (ec2_sg_id)"),
        ("RDS", "../../modules/rds", "VPC (subnet_ids), SG (rds_sg_id)"),
        ("ECR", "../../modules/ecr", "없음"),
    ]
    for i, data in enumerate(dep_data):
        add_data_row(dep_table, i + 1, list(data))

    doc.add_paragraph()
    add_heading_styled(doc, "environments/dev/main.tf", level=3)
    dev_main = read_tf_file("environments/dev/main.tf")
    add_code_block(doc, dev_main)

    doc.add_paragraph()
    add_heading_styled(doc, "environments/dev/variables.tf", level=3)
    dev_vars = read_tf_file("environments/dev/variables.tf")
    add_code_block(doc, dev_vars)

    doc.add_paragraph()
    add_heading_styled(doc, "environments/dev/output.tf", level=3)
    dev_outputs = read_tf_file("environments/dev/output.tf")
    add_code_block(doc, dev_outputs)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 4. 변수 명세
    # ════════════════════════════════════════
    add_heading_styled(doc, "4. 변수 명세", level=1)
    doc.add_paragraph(
        "아래 표는 전체 프로젝트에서 사용되는 변수를 모듈별로 정리한 것입니다."
    )
    doc.add_paragraph()

    # 환경 변수 (최상위)
    p = doc.add_paragraph()
    run = p.add_run("[환경 변수 (environments/dev)]")
    run.bold = True
    run.font.size = Pt(10)

    env_var_table = doc.add_table(rows=7, cols=5)
    env_var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(env_var_table, 0, ["변수명", "타입", "기본값", "필수", "설명"])
    env_vars_data = [
        ("aws_region", "string", "ap-northeast-2", "X", "AWS 리전"),
        ("project_name", "string", "-", "O", "프로젝트 이름"),
        ("environment", "string", "-", "O", "환경 (dev, prod)"),
        ("vpc_cidr", "string", "10.0.0.0/16", "X", "VPC CIDR 블록"),
        ("key_name", "string", "-", "O", "EC2 SSH 키페어 이름"),
        ("db_password", "string (sensitive)", "-", "O", "RDS 비밀번호"),
    ]
    for i, data in enumerate(env_vars_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(env_var_table, i + 1, list(data),
                     [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # VPC 모듈 변수
    p = doc.add_paragraph()
    run = p.add_run("[VPC 모듈 변수]")
    run.bold = True
    run.font.size = Pt(10)

    vpc_var_table = doc.add_table(rows=9, cols=5)
    vpc_var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(vpc_var_table, 0, ["변수명", "타입", "기본값", "필수", "설명"])
    vpc_vars_data = [
        ("project_name", "string", "-", "O", "프로젝트 이름"),
        ("environment", "string", "-", "O", "환경 (dev, prod)"),
        ("vpc_cidr", "string", "10.0.0.0/16", "X", "VPC CIDR 블록"),
        ("public_subnet_cidr", "string", "10.0.1.0/24", "X", "퍼블릭 서브넷 CIDR"),
        ("private_subnet_cidr", "string", "10.0.10.0/24", "X", "프라이빗 서브넷 CIDR"),
        ("availability_zone", "string", "ap-northeast-2a", "X", "가용영역"),
        ("availability_zone_2", "string", "ap-northeast-2c", "X", "두 번째 가용영역"),
        ("private_subnet_cidr_2", "string", "10.0.20.0/24", "X", "두 번째 프라이빗 서브넷 CIDR"),
    ]
    for i, data in enumerate(vpc_vars_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(vpc_var_table, i + 1, list(data),
                     [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # EC2 모듈 변수
    p = doc.add_paragraph()
    run = p.add_run("[EC2 모듈 변수]")
    run.bold = True
    run.font.size = Pt(10)

    ec2_var_table = doc.add_table(rows=7, cols=5)
    ec2_var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(ec2_var_table, 0, ["변수명", "타입", "기본값", "필수", "설명"])
    ec2_vars_data = [
        ("project_name", "string", "-", "O", "프로젝트 이름"),
        ("environment", "string", "-", "O", "환경 (dev, prod)"),
        ("subnet_id", "string", "-", "O", "EC2 서브넷 ID"),
        ("security_group_id", "string", "-", "O", "EC2 보안 그룹 ID"),
        ("instance_type", "string", "t3.micro", "X", "EC2 인스턴스 타입"),
        ("key_name", "string", "-", "O", "SSH 키페어 이름"),
    ]
    for i, data in enumerate(ec2_vars_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(ec2_var_table, i + 1, list(data),
                     [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # RDS 모듈 변수
    p = doc.add_paragraph()
    run = p.add_run("[RDS 모듈 변수]")
    run.bold = True
    run.font.size = Pt(10)

    rds_var_table = doc.add_table(rows=8, cols=5)
    rds_var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(rds_var_table, 0, ["변수명", "타입", "기본값", "필수", "설명"])
    rds_vars_data = [
        ("project_name", "string", "-", "O", "프로젝트 이름"),
        ("environment", "string", "-", "O", "환경 (dev, prod)"),
        ("subnet_ids", "list(string)", "-", "O", "RDS 서브넷 ID 목록"),
        ("security_group_id", "string", "-", "O", "RDS 보안 그룹 ID"),
        ("db_name", "string", "finops", "X", "데이터베이스 이름"),
        ("db_username", "string", "postgres", "X", "데이터베이스 사용자명"),
        ("db_password", "string (sensitive)", "-", "O", "데이터베이스 비밀번호"),
    ]
    for i, data in enumerate(rds_vars_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(rds_var_table, i + 1, list(data),
                     [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph()

    # ECR 모듈 변수
    p = doc.add_paragraph()
    run = p.add_run("[ECR 모듈 변수]")
    run.bold = True
    run.font.size = Pt(10)

    ecr_var_table = doc.add_table(rows=4, cols=5)
    ecr_var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(ecr_var_table, 0, ["변수명", "타입", "기본값", "필수", "설명"])
    ecr_vars_data = [
        ("project_name", "string", "-", "O", "프로젝트 이름"),
        ("environment", "string", "-", "O", "환경 (dev, prod)"),
        ("repository_names", "list(string)", "[5개 서비스]", "X", "ECR 레포지토리 이름 목록"),
    ]
    for i, data in enumerate(ecr_vars_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(ecr_var_table, i + 1, list(data),
                     [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_page_break()

    # ════════════════════════════════════════
    # 5. 출력값 명세
    # ════════════════════════════════════════
    add_heading_styled(doc, "5. 출력값 명세", level=1)
    doc.add_paragraph(
        "각 모듈 및 환경 구성에서 출력되는 값을 정리합니다."
    )
    doc.add_paragraph()

    outputs_table = doc.add_table(rows=15, cols=4)
    outputs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(outputs_table, 0, ["모듈", "출력명", "설명", "참조 예시"])
    outputs_data = [
        ("VPC", "vpc_id", "VPC ID", "module.vpc.vpc_id"),
        ("VPC", "public_subnet_id", "퍼블릭 서브넷 ID", "module.vpc.public_subnet_id"),
        ("VPC", "private_subnet_id", "프라이빗 서브넷 ID", "module.vpc.private_subnet_id"),
        ("VPC", "private_subnet_id_2", "두 번째 프라이빗 서브넷 ID", "module.vpc.private_subnet_id_2"),
        ("Security Group", "ec2_security_group_id", "EC2 보안 그룹 ID", "module.security_group.ec2_security_group_id"),
        ("Security Group", "rds_security_group_id", "RDS 보안 그룹 ID", "module.security_group.rds_security_group_id"),
        ("EC2", "instance_id", "EC2 인스턴스 ID", "module.ec2.instance_id"),
        ("EC2", "public_ip", "EC2 퍼블릭 IP", "module.ec2.public_ip"),
        ("EC2", "public_dns", "EC2 퍼블릭 DNS", "module.ec2.public_dns"),
        ("ECR", "repository_urls", "ECR 레포지토리 URL 맵", "module.ecr.repository_urls"),
        ("ECR", "repository_arns", "ECR 레포지토리 ARN 맵", "module.ecr.repository_arns"),
        ("RDS", "endpoint", "RDS 엔드포인트 (host:port)", "module.rds.endpoint"),
        ("RDS", "address", "RDS 주소 (포트 제외)", "module.rds.address"),
        ("RDS", "port", "RDS 포트 (5432)", "module.rds.port"),
    ]
    for i, data in enumerate(outputs_data):
        add_data_row(outputs_table, i + 1, list(data))

    doc.add_page_break()

    # ════════════════════════════════════════
    # 6. 리소스 요약
    # ════════════════════════════════════════
    add_heading_styled(doc, "6. 리소스 요약", level=1)
    doc.add_paragraph(
        "본 Terraform 프로젝트를 통해 프로비저닝되는 전체 AWS 리소스를 요약합니다."
    )
    doc.add_paragraph()

    summary_table = doc.add_table(rows=16, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(summary_table, 0, ["No", "AWS 서비스", "리소스 타입", "수량"])
    summary_data = [
        ("1", "VPC", "aws_vpc", "1"),
        ("2", "Internet Gateway", "aws_internet_gateway", "1"),
        ("3", "Subnet (Public)", "aws_subnet", "1"),
        ("4", "Subnet (Private)", "aws_subnet", "2"),
        ("5", "Route Table", "aws_route_table", "1"),
        ("6", "Route Table Association", "aws_route_table_association", "1"),
        ("7", "Security Group (EC2)", "aws_security_group", "1"),
        ("8", "Security Group (RDS)", "aws_security_group", "1"),
        ("9", "IAM Role", "aws_iam_role", "1"),
        ("10", "IAM Instance Profile", "aws_iam_instance_profile", "1"),
        ("11", "IAM Policy Attachment", "aws_iam_role_policy_attachment", "1"),
        ("12", "EC2 Instance", "aws_instance", "1"),
        ("13", "ECR Repository", "aws_ecr_repository", "5"),
        ("14", "ECR Lifecycle Policy", "aws_ecr_lifecycle_policy", "5"),
        ("15", "RDS Instance", "aws_db_instance", "1"),
    ]
    for i, data in enumerate(summary_data):
        center = WD_ALIGN_PARAGRAPH.CENTER
        add_data_row(summary_table, i + 1, list(data),
                     [center, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, center])

    doc.add_paragraph()

    # 총계
    p = doc.add_paragraph()
    run = p.add_run("총 생성 리소스: 24개")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_paragraph()

    # 비용 관련 참고사항
    p = doc.add_paragraph()
    run = p.add_run("[비용 최적화 참고사항]")
    run.bold = True
    run.font.size = Pt(10)

    cost_items = [
        "EC2 인스턴스: t3.micro (프리티어 적용 가능)",
        "RDS 인스턴스: db.t3.micro (프리티어 적용 가능)",
        "ECR: 이미지 수명주기 정책으로 최근 5개만 보관 (스토리지 비용 절감)",
        "RDS 백업: 보존 기간 0일 (백업 비용 절감, 개발 환경)",
        "Multi-AZ / 암호화: 비활성화 (개발 환경 비용 절감)",
    ]
    for item in cost_items:
        p = doc.add_paragraph(f"  - {item}")
        for run in p.runs:
            run.font.size = Pt(9)

    # ── 저장 ──
    output_path = os.path.join(BASE_DIR, "FinOps_인프라_코드_산출물.docx")
    doc.save(output_path)
    print(f"문서가 생성되었습니다: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
